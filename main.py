"""res-sum Streamlit App — Interactive research paper summarization with GraphRAG."""

import os
import tempfile
import json
from io import BytesIO
from pathlib import Path

import streamlit as st

st.set_page_config(
    page_title="res-sum — Research Evidence Synthesis",
    page_icon="📚",
    layout="wide",
    initial_sidebar_state="expanded",
)

# ---------------------------------------------------------------------------
# Session state initialization
# ---------------------------------------------------------------------------
if "rs" not in st.session_state:
    st.session_state.rs = None
if "ingested" not in st.session_state:
    st.session_state.ingested = False
if "summaries" not in st.session_state:
    st.session_state.summaries = {}
if "dashboard_html" not in st.session_state:
    st.session_state.dashboard_html = None
if "temp_dir" not in st.session_state:
    st.session_state.temp_dir = tempfile.mkdtemp(prefix="res_sum_")
if "stop_requested" not in st.session_state:
    st.session_state.stop_requested = False
if "error_log" not in st.session_state:
    st.session_state.error_log = []


# ---------------------------------------------------------------------------
# Sidebar — Configuration
# ---------------------------------------------------------------------------
st.sidebar.image(
    "https://img.shields.io/pypi/v/res-sum?style=flat-square&label=res-sum",
    width=150,
)
st.sidebar.title("Configuration")

# Provider selection
provider = st.sidebar.selectbox(
    "LLM Provider",
    options=["ollama_cloud", "ollama", "groq", "openai", "anthropic"],
    index=0,
    help="Choose your LLM provider. Ollama Cloud is free with an API key.",
)

# Model input
default_models = {
    "ollama_cloud": "gpt-oss:20b-cloud",
    "ollama": "llama3.2",
    "groq": "llama-3.3-70b-versatile",
    "openai": "gpt-4o-mini",
    "anthropic": "claude-sonnet-4-20250514",
}
model = st.sidebar.text_input(
    "Model",
    value=default_models.get(provider, "llama3.2"),
    help="Model name for your chosen provider.",
)

# API Key
api_key = st.sidebar.text_input(
    "API Key",
    type="password",
    help="Required for cloud providers. Set as env var to avoid re-entering.",
    value=os.environ.get({
        "ollama_cloud": "OLLAMA_API_KEY",
        "groq": "GROQ_API_KEY",
        "openai": "OPENAI_API_KEY",
        "anthropic": "ANTHROPIC_API_KEY",
    }.get(provider, "OLLAMA_API_KEY"), ""),
)

# Domain
domain = st.sidebar.selectbox(
    "Domain",
    options=["ecology", "general"],
    index=0,
    help="Domain-specific prompts and entity schemas.",
)

# Start section
start_section = st.sidebar.selectbox(
    "Start Summary From",
    options=["methodology", "introduction"],
    index=0,
    help="Which section to start extracting text for summaries. Knowledge graph always uses full paper.",
)

# Retrieval mode
retrieval_mode = st.sidebar.selectbox(
    "Retrieval Mode",
    options=["hybrid", "local", "graph", "global"],
    index=0,
    help="hybrid = vector + graph + community (recommended). local = vector only. graph = graph traversal. global = community summaries.",
)

# Output format
output_format = st.sidebar.selectbox(
    "Output Format",
    options=["docx", "json", "csv"],
    index=0,
)

# Advanced settings
with st.sidebar.expander("Advanced Settings"):
    build_graph = st.checkbox("Build Knowledge Graph", value=True, help="Extract entities and relationships into a knowledge graph.")
    use_cot = st.checkbox("Chain-of-Thought Prompting", value=True, help="Use structured CoT prompts for better extraction.")
    n_chunks = st.slider("Chunks to retrieve", min_value=3, max_value=30, value=10, help="Number of chunks to use as context for summarization.")


# ---------------------------------------------------------------------------
# Initialize ResSum
# ---------------------------------------------------------------------------
def init_res_sum():
    """Initialize or reinitialize the ResSum instance."""
    from res_sum import ResSum

    data_dir = os.path.join(st.session_state.temp_dir, "knowledge_base")

    kwargs = {
        "llm_provider": provider,
        "model": model,
        "domain": domain,
        "data_dir": data_dir,
        "build_graph": build_graph,
        "start_section": start_section,
    }
    if api_key:
        kwargs["api_key"] = api_key

    st.session_state.rs = ResSum(**kwargs)
    return st.session_state.rs


# ---------------------------------------------------------------------------
# Main content — Tabs
# ---------------------------------------------------------------------------
st.title("📚 res-sum — Research Evidence Synthesis")
st.caption("Upload research papers, build knowledge graphs, and generate structured summaries using LLMs.")

tab_analysis, tab_kb, tab_vectorstore = st.tabs([
    "🔬 Analysis",
    "🔗 Knowledge Base",
    "🔍 Semantic Search",
])

# ---------------------------------------------------------------------------
# Tab 1: Analysis (combined ingest + summarize)
# ---------------------------------------------------------------------------
with tab_analysis:
    st.header("Upload & Analyze Papers")

    # File uploader
    uploaded_files = st.file_uploader(
        "Upload PDF files",
        type=["pdf"],
        accept_multiple_files=True,
        help="Upload one or more research papers in PDF format.",
    )

    if uploaded_files:
        st.success(f"{len(uploaded_files)} file(s) uploaded.")

    # Custom prompt
    prompt = st.text_area(
        "What would you like to extract or synthesize from these papers?",
        value="Provide a comprehensive summary of this research paper, focusing on the research problem, methods used, key findings, and their implications.",
        height=100,
        help="This prompt guides both knowledge graph extraction and the final synthesis.",
    )

    col1, col2 = st.columns([3, 1])

    with col1:
        run_btn = st.button("🔬 Run Analysis", use_container_width=True, type="primary")

    with col2:
        stop_btn = st.button("⛔ Stop", use_container_width=True)

    if stop_btn:
        st.session_state.stop_requested = True
        st.warning("Stop requested. The current operation will halt after the current paper.")

    if run_btn:
        st.session_state.stop_requested = False
        st.session_state.error_log = []

        if not uploaded_files:
            st.error("Please upload at least one PDF file.")
        elif not api_key and provider != "ollama":
            st.error(f"Please provide an API key for {provider}.")
        else:
            # Save uploaded files to temp directory
            pdf_dir = os.path.join(st.session_state.temp_dir, "pdfs")
            os.makedirs(pdf_dir, exist_ok=True)

            for f in uploaded_files:
                filepath = os.path.join(pdf_dir, f.name)
                with open(filepath, "wb") as out:
                    out.write(f.getbuffer())

            # --- Phase 1: Ingest ---
            status = st.status("**Phase 1/2:** Ingesting papers — building knowledge base...", expanded=True)

            with status:
                try:
                    rs = init_res_sum()
                    results = rs.ingest_papers(pdf_dir, rebuild=True)
                    st.session_state.ingested = True

                    for r in results:
                        if st.session_state.stop_requested:
                            st.warning("Stopped by user.")
                            break
                        if r.ingested:
                            st.write(f"✅ {r.filename}: {r.num_chunks} chunks, {r.num_entities} entities")
                        else:
                            st.write(f"⚠️ {r.filename}: failed")
                            st.session_state.error_log.append(f"Ingestion failed: {r.filename}")

                    # Generate dashboard HTML for other tabs
                    dashboard_path = rs.explore(
                        output_path=os.path.join(st.session_state.temp_dir, "dashboard.html"),
                        open_browser=False,
                    )
                    st.session_state.dashboard_html = Path(dashboard_path).read_text(encoding="utf-8")

                    stats = rs.stats()
                    st.write(
                        f"Knowledge base: **{stats['graph_nodes']}** nodes, "
                        f"**{stats['graph_edges']}** edges, "
                        f"**{stats['total_chunks']}** chunks"
                    )
                    status.update(label="**Phase 1/2:** Ingestion complete", state="complete")

                except Exception as e:
                    st.session_state.error_log.append(f"Ingestion error: {e}")
                    status.update(label="**Phase 1/2:** Ingestion failed", state="error")
                    st.error(f"Ingestion failed: {e}")

            # --- Phase 2: Summarize ---
            if st.session_state.ingested and not st.session_state.stop_requested:
                rs = st.session_state.rs
                papers = rs.vector_store.list_papers() if hasattr(rs.vector_store, "list_papers") else []

                status2 = st.status(f"**Phase 2/2:** Generating summaries for {len(papers)} paper(s)...", expanded=True)

                with status2:
                    summaries = {}
                    for idx, paper in enumerate(papers):
                        if st.session_state.stop_requested:
                            st.warning("Stopped by user.")
                            break
                        try:
                            st.write(f"Processing {paper}...")
                            summary = rs.summarize(
                                query=prompt,
                                prompt=prompt,
                                use_cot=use_cot,
                                n_chunks=n_chunks,
                                paper_filter=paper,
                                mode=retrieval_mode,
                            )
                            summaries[paper] = summary
                            st.write(f"✅ {paper}")
                        except Exception as e:
                            st.write(f"❌ {paper}: {e}")
                            st.session_state.error_log.append(f"Summary failed for {paper}: {e}")

                    st.session_state.summaries = summaries
                    status2.update(label=f"**Phase 2/2:** {len(summaries)} summary(ies) generated", state="complete")

            # Show error log if any
            if st.session_state.error_log:
                with st.expander("⚠️ Error Log", expanded=False):
                    for err in st.session_state.error_log:
                        st.code(err, language=None)

    # --- Display & Download Summaries ---
    if st.session_state.summaries:
        st.markdown("---")
        st.header("Generated Summaries")

        for paper, summary in st.session_state.summaries.items():
            with st.expander(f"📄 {paper}", expanded=True):
                st.markdown(summary)

                # Download buttons
                clean_name = os.path.splitext(paper)[0]

                if output_format == "docx":
                    try:
                        from docx import Document
                        doc = Document()
                        doc.add_heading(f"Summary — {clean_name}", level=1)
                        doc.add_paragraph(summary)
                        buf = BytesIO()
                        doc.save(buf)
                        buf.seek(0)
                        st.download_button(
                            f"Download DOCX",
                            data=buf,
                            file_name=f"Summary-{clean_name}.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            key=f"dl_docx_{paper}",
                        )
                    except ImportError:
                        st.warning("python-docx not installed. Install with: pip install python-docx")

                elif output_format == "json":
                    data = json.dumps({"title": paper, "summary": summary}, indent=2)
                    st.download_button(
                        f"Download JSON",
                        data=data,
                        file_name=f"Summary-{clean_name}.json",
                        mime="application/json",
                        key=f"dl_json_{paper}",
                    )

                elif output_format == "csv":
                    import csv
                    buf = BytesIO()
                    writer = csv.writer(buf := BytesIO())
                    buf.write(b"title,summary\n")
                    # Escape CSV properly
                    csv_line = f'"{paper}","{summary.replace(chr(34), chr(34)+chr(34))}"\n'
                    buf.write(csv_line.encode())
                    buf.seek(0)
                    st.download_button(
                        f"Download CSV",
                        data=buf,
                        file_name=f"Summary-{clean_name}.csv",
                        mime="text/csv",
                        key=f"dl_csv_{paper}",
                    )

        # Download all as ZIP
        st.markdown("---")
        if st.button("📥 Download All Summaries as ZIP"):
            import zipfile
            zip_buf = BytesIO()
            with zipfile.ZipFile(zip_buf, "w", zipfile.ZIP_DEFLATED) as zf:
                for paper, summary in st.session_state.summaries.items():
                    clean_name = os.path.splitext(paper)[0]
                    zf.writestr(f"Summary-{clean_name}.txt", summary)
            zip_buf.seek(0)
            st.download_button(
                "📥 Download ZIP",
                data=zip_buf,
                file_name="all_summaries.zip",
                mime="application/zip",
                key="dl_zip",
            )


# ---------------------------------------------------------------------------
# Tab 2: Explore Knowledge Base (embedded dashboard)
# ---------------------------------------------------------------------------
# Tab 2: Knowledge Base (interactive dashboard)
# ---------------------------------------------------------------------------
with tab_kb:
    st.header("Knowledge Base")
    st.caption("Interactive dashboard with Overview, Knowledge Graph, Vector Store preview, and Communities.")

    if st.session_state.dashboard_html:
        import streamlit.components.v1 as components
        components.html(st.session_state.dashboard_html, height=1200, scrolling=True)
    else:
        st.info("Run analysis first to build the knowledge base. Go to the **Analysis** tab and click **Run Analysis**.")


# ---------------------------------------------------------------------------
# Tab 3: Vector Store (live ChromaDB search + browse)
# ---------------------------------------------------------------------------
with tab_vectorstore:
    st.header("Semantic Search")

    if st.session_state.rs and st.session_state.ingested:
        rs = st.session_state.rs
        vs = rs.vector_store

        search_subtab, browse_subtab = st.tabs(["🔍 Search (Cosine Similarity)", "📄 Browse Chunks"])

        with search_subtab:
            st.caption("Search uses real cosine similarity on ChromaDB embeddings — the same retrieval used during analysis.")

            papers = vs.list_papers() if hasattr(vs, "list_papers") else []
            search_paper = st.selectbox(
                "Filter by paper",
                ["All papers"] + papers,
                key="search_paper_filter",
            )

            search_query = st.text_input(
                "Search query",
                placeholder="e.g., species distribution modeling, trophic cascade effects...",
                key="vector_search_query",
            )
            search_n = st.slider("Number of results", min_value=3, max_value=30, value=10, key="vector_search_n")

            if search_query:
                where = {"paper_filename": search_paper} if search_paper != "All papers" else None
                results = vs.search(search_query, n_results=search_n, where=where)
                docs = results.get("documents", [[]])[0]
                ids = results.get("ids", [[]])[0]
                distances = results.get("distances", [[]])[0]
                metadatas = results.get("metadatas", [[]])[0]

                st.markdown(f"**{len(docs)} results** for *\"{search_query}\"*")

                for rank, (doc, doc_id, dist, meta) in enumerate(zip(docs, ids, distances, metadatas), 1):
                    score = 1.0 / (1.0 + dist)
                    paper = meta.get("paper_filename", "unknown") if meta else "unknown"
                    section = meta.get("section", "unknown") if meta else "unknown"

                    if score > 0.6:
                        color = "🟢"
                    elif score > 0.4:
                        color = "🟡"
                    else:
                        color = "🔴"

                    with st.expander(f"{color} #{rank} — Score: **{score:.4f}** — {paper} ({section})"):
                        st.markdown(f"**Cosine Similarity Score:** `{score:.6f}`")
                        st.markdown(f"**Paper:** {paper}")
                        st.markdown(f"**Section:** {section}")
                        st.markdown(f"**Chunk ID:** `{doc_id}`")
                        st.markdown("---")
                        st.markdown(doc)
            else:
                st.info("Enter a query above to search with cosine similarity.")

        with browse_subtab:
            st.caption("Browse all chunks stored in the vector database.")

            papers = vs.list_papers() if hasattr(vs, "list_papers") else []
            total = vs.count()
            st.metric("Total Chunks", total)

            if papers:
                selected_paper = st.selectbox(
                    "Filter by paper",
                    ["All papers"] + papers,
                    key="browse_paper_filter",
                )
                where = {"paper_filename": selected_paper} if selected_paper != "All papers" else None

                try:
                    if where:
                        all_data = vs.collection.get(where=where, include=["metadatas", "documents"])
                    else:
                        all_data = vs.collection.get(include=["metadatas", "documents"])

                    num_chunks = len(all_data.get("ids", []))
                    st.markdown(f"Showing **{min(num_chunks, 50)}** of **{num_chunks}** chunks")

                    for i in range(min(num_chunks, 50)):
                        doc = all_data["documents"][i] if all_data.get("documents") else ""
                        meta = all_data["metadatas"][i] if all_data.get("metadatas") else {}
                        section = meta.get("section", "unknown")
                        chunk_idx = meta.get("chunk_index", i)

                        with st.expander(f"Chunk {chunk_idx + 1} — Section: {section} — `{all_data['ids'][i][:12]}...`"):
                            st.markdown(doc)
                except Exception as e:
                    st.error(f"Failed to browse chunks: {e}")
    else:
        st.info("Run analysis first. Go to the **Analysis** tab and click **Run Analysis**.")


# ---------------------------------------------------------------------------
# Footer
# ---------------------------------------------------------------------------
st.markdown("---")
st.markdown(
    '<p style="text-align:center;color:#888;font-size:13px;">'
    'Powered by <a href="https://pypi.org/project/res-sum/" target="_blank">res-sum</a> '
    '— Open-source research evidence synthesis with GraphRAG'
    '</p>',
    unsafe_allow_html=True,
)
